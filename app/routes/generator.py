import os
from flask import Blueprint, render_template, jsonify, request, current_app, send_file
from docx import Document as DocxDocument
from app.utils.json import JsonUtils
from app.utils.document import Document as DocumentUtils


generator_blueprint = Blueprint("generator", __name__, url_prefix="/")


@generator_blueprint.route("/")
def index():
    return render_template("index.html")


@generator_blueprint.route("/api/state", methods=["GET"])
def get_state():
    return jsonify(JsonUtils.load_json())


@generator_blueprint.route("/api/save-variable", methods=["POST"])
def save_variable():
    payload = request.get_json(force=True)
    tab = payload.get("tab")
    key = payload.get("key")
    value = payload.get("value", "")
    old_key = payload.get("old_key")

    state = JsonUtils.load_json()

    if tab not in state["tabs"]:
        return jsonify({"ok": False, "error": "Aba inválida."}), 400

    if tab not in state["data"]:
        state["data"][tab] = {}

    if old_key and old_key != key and old_key in state["data"][tab]:
        del state["data"][tab][old_key]

    state["data"][tab][key] = value
    JsonUtils.save_json(state)
    return jsonify({"ok": True, "state": state})


@generator_blueprint.route("/api/delete-variable", methods=["POST"])
def delete_variable():
    payload = request.get_json(force=True)
    tab = payload.get("tab")
    key = payload.get("key")

    state = JsonUtils.load_json()

    if tab in state["data"] and key in state["data"][tab]:
        del state["data"][tab][key]
        JsonUtils.save_json(state)

    return jsonify({"ok": True, "state": state})


@generator_blueprint.route("/api/preview-docx", methods=["POST"])
def preview_docx():
    if "file" not in request.files:
        return jsonify({"ok": False, "error": "Arquivo DOCX não enviado."}), 400

    file = request.files["file"]

    if not file or not file.filename:
        return jsonify({"ok": False, "error": "Nenhum arquivo selecionado."}), 400

    if not file.filename.lower().endswith(".docx"):
        return jsonify({"ok": False, "error": "Envie um arquivo .docx"}), 400

    saved_path = os.path.join(current_app.config["UPLOAD_DIR"], "modelo.docx")
    file.save(saved_path)

    preview_lines = []

    try:
        doc = DocxDocument(saved_path)

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                preview_lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    preview_lines.append(row_text)

        preview_text = "\n".join(preview_lines[:300])

        return jsonify({
            "ok": True,
            "preview": preview_text or "Documento lido, mas sem texto visível."
        })
    except Exception as e:
        return jsonify({"ok": False, "error": f"Erro ao ler DOCX: {str(e)}"}), 500


@generator_blueprint.route("/api/generate-docx", methods=["POST"])
def generate_docx():
    source_path = os.path.join(current_app.config["UPLOAD_DIR"], "modelo.docx")
    if not os.path.exists(source_path):
        return jsonify({
            "ok": False,
            "error": "Nenhum documento modelo foi carregado. Clique em 'Selecionar / Ler documento' antes."
        }), 400

    state = JsonUtils.load_json()
    replacements = DocumentUtils.flatten_variables(state)

    output_path = os.path.join(current_app.config["UPLOAD_DIR"], "documento_preenchido.docx")
    DocumentUtils.replace_in_docx(source_path, output_path, replacements)

    return send_file(
        output_path,
        as_attachment=True,
        download_name="documento_preenchido.docx"
    )

@generator_blueprint.route("/version-details", methods=["GET"])
def version_details():
    return render_template("version_details.html")
  