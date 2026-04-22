from docx import Document as DocxDocument


class Document:
    @staticmethod
    def replace_in_paragraph(paragraph, replacements):
        full_text = "".join(run.text for run in paragraph.runs)

        if not full_text:
            return

        updated = full_text
        for key, value in replacements.items():
            updated = updated.replace(key, value if value is not None else "")

        if updated != full_text:
            if paragraph.runs:
                paragraph.runs[0].text = updated
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = updated

    @staticmethod
    def replace_in_docx(input_path, output_path, replacements):
        doc = DocxDocument(input_path)

        for paragraph in doc.paragraphs:
            Document.replace_in_paragraph(paragraph, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        Document.replace_in_paragraph(paragraph, replacements)

        doc.save(output_path)

    @staticmethod
    def flatten_variables(data_obj):
        merged = {}
        tabs = data_obj.get("tabs", [])
        tab_data = data_obj.get("data", {})

        for tab in tabs:
            for key, value in tab_data.get(tab, {}).items():
                merged[key] = value

        return merged
    